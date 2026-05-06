#!/usr/bin/env python3
"""将配置中 md_to_docx.input_dir 下所有 .md 转为 .docx 写入 output_dir（见 config/config.yaml）。
依赖: pip install python-docx PyYAML
"""
from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.shared import Pt

_REPO = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(_REPO))
from config.paths import REPO_ROOT, md_to_docx_dirs


def _add_inline_runs(paragraph, text: str) -> None:
    """解析段落内 **粗体**、*斜体*、`行内代码`（不支持嵌套交叉）。"""
    i = 0
    n = len(text)
    while i < n:
        if text.startswith("**", i):
            end = text.find("**", i + 2)
            if end == -1:
                paragraph.add_run(text[i:])
                break
            run = paragraph.add_run(text[i + 2 : end])
            run.bold = True
            i = end + 2
            continue
        if text.startswith("`", i):
            end = text.find("`", i + 1)
            if end == -1:
                paragraph.add_run(text[i:])
                break
            run = paragraph.add_run(text[i + 1 : end])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            i = end + 1
            continue
        if text[i] == "*" and (i + 1 >= n or text[i + 1] != "*"):
            end = text.find("*", i + 1)
            if end == -1:
                paragraph.add_run(text[i:])
                break
            run = paragraph.add_run(text[i + 1 : end])
            run.italic = True
            i = end + 1
            continue
        nxt = n
        for token in ("**", "`", "*"):
            j = text.find(token, i + 1)
            if j != -1 and j < nxt:
                nxt = j
        paragraph.add_run(text[i:nxt])
        i = nxt


def _add_plain_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _add_inline_runs(p, text)


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return "|" in s and s.startswith("|")


def _is_table_sep(line: str) -> bool:
    s = line.strip().strip("|")
    if not s:
        return False
    parts = [p.strip() for p in s.split("|")]
    return all(re.fullmatch(r":?-{3,}:?", p) for p in parts if p)


def _parse_table_row(line: str) -> list[str]:
    row = [c.strip() for c in line.strip().strip("|").split("|")]
    return row


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    w = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=w)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(w):
            cell_text = row[ci] if ci < len(row) else ""
            table.rows[ri].cells[ci].text = cell_text


def _iter_markdown_blocks(lines: list[str]) -> Iterator[tuple[str, object]]:
    """产出 (kind, payload)。kind: heading, para, bullet, ordered, code, table, hr。"""
    i = 0
    n = len(lines)

    def flush_para(buf: list[str]) -> None:
        if not buf:
            return
        text = " ".join(x.rstrip() for x in buf).strip()
        if text:
            yield ("para", text)
        buf.clear()

    para_buf: list[str] = []

    while i < n:
        raw = lines[i]
        line = raw.rstrip("\n")
        stripped = line.strip()

        if not stripped:
            yield from flush_para(para_buf)
            i += 1
            continue

        if stripped == "---" or stripped == "***" or stripped == "___":
            yield from flush_para(para_buf)
            yield ("hr", None)
            i += 1
            continue

        if stripped.startswith("```"):
            yield from flush_para(para_buf)
            lang = stripped[3:].strip()
            code_lines: list[str] = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            if i < n:
                i += 1
            yield ("code", (lang, "\n".join(code_lines)))
            continue

        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            yield from flush_para(para_buf)
            yield ("heading", (len(m.group(1)), m.group(2).strip()))
            i += 1
            continue

        if re.match(r"^[-*+]\s+", stripped):
            yield from flush_para(para_buf)
            yield ("bullet", re.sub(r"^[-*+]\s+", "", stripped))
            i += 1
            continue

        om = re.match(r"^(\d+)\.\s+", stripped)
        if om:
            yield from flush_para(para_buf)
            yield ("ordered", stripped[om.end() :].strip())
            i += 1
            continue

        if _is_table_row(line) and i + 1 < n and _is_table_sep(lines[i + 1]):
            yield from flush_para(para_buf)
            rows: list[list[str]] = [_parse_table_row(line)]
            i += 1
            if i < n:
                i += 1
            while i < n and _is_table_row(lines[i]):
                rows.append(_parse_table_row(lines[i]))
                i += 1
            yield ("table", rows)
            continue

        para_buf.append(stripped)
        i += 1

    yield from flush_para(para_buf)


def markdown_to_docx(md_text: str) -> Document:
    doc = Document()
    lines = md_text.splitlines()
    for kind, payload in _iter_markdown_blocks(lines):
        if kind == "heading":
            level, text = payload  # type: ignore[misc]
            level = min(max(int(level), 1), 9)
            p = doc.add_paragraph(style=f"Heading {level}")
            _add_inline_runs(p, text)
        elif kind == "para":
            _add_plain_paragraph(doc, str(payload))
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, str(payload))
        elif kind == "ordered":
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, str(payload))
        elif kind == "code":
            _lang, body = payload  # type: ignore[misc]
            p = doc.add_paragraph()
            run = p.add_run(body)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif kind == "table":
            _add_table(doc, payload)  # type: ignore[arg-type]
        elif kind == "hr":
            doc.add_paragraph("---")
    return doc


def main() -> None:
    input_dir, output_dir = md_to_docx_dirs()
    if not input_dir.is_dir():
        sys.exit(f"输入目录不存在: {input_dir}")
    output_dir.mkdir(parents=True, exist_ok=True)
    md_files = sorted(input_dir.glob("*.md"))
    if not md_files:
        print(f"未在 {input_dir} 找到 .md 文件")
        return
    for src in md_files:
        dest = output_dir / f"{src.stem}_python_docx.docx"
        try:
            text = src.read_text(encoding="utf-8")
            doc = markdown_to_docx(text)
            doc.save(dest)
            print(f"完成: {src.name} → {dest.relative_to(REPO_ROOT)}")
        except Exception as e:
            print(f"跳过: {src.name} ({e})")


if __name__ == "__main__":
    main()
