#!/usr/bin/env python3
"""将配置中 docx_to_md.input_dir 下所有 .docx 转为 Markdown 写入 output_dir（见 config/config.yaml）。
依赖: pip install python-docx PyYAML
"""
import sys
from pathlib import Path
from typing import Optional

from docx import Document


def _heading_level(style_name: str) -> Optional[int]:
    if style_name.startswith("Heading "):
        try:
            return int(style_name.split()[1])
        except (IndexError, ValueError):
            return None
    if style_name == "Title":
        return 1
    if style_name == "Subtitle":
        return 2
    return None


def _runs_to_md(para) -> str:
    out: list[str] = []
    for run in para.runs:
        t = run.text
        if not t:
            continue
        if run.bold and run.italic:
            t = f"***{t}***"
        elif run.bold:
            t = f"**{t}**"
        elif run.italic:
            t = f"*{t}*"
        out.append(t)
    return "".join(out)


def _paragraph_to_md(para) -> Optional[str]:
    text = _runs_to_md(para).strip() or para.text.strip()
    if not text:
        return None
    style = para.style.name
    hl = _heading_level(style)
    if hl:
        return f"{'#' * hl} {text}"
    if "List" in style:
        return f"- {text}"
    return text


def _tables_to_md(doc) -> str:
    blocks: list[str] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([c.text.replace("\n", " ").replace("|", "\\|") for c in row.cells])
        if not rows:
            continue
        w = len(rows[0])
        sep = "| " + " | ".join(["---"] * w) + " |"
        lines = ["| " + " | ".join(r) + " |" for r in rows]
        lines.insert(1, sep)
        blocks.append("\n".join(lines))
    return "\n\n".join(blocks)


def docx_to_markdown(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        line = _paragraph_to_md(para)
        if line:
            parts.append(line)
    body = "\n\n".join(parts)
    tbl = _tables_to_md(doc)
    if tbl:
        body = f"{body}\n\n{tbl}" if body else tbl
    return body


def main() -> None:
    _repo = Path(__file__).resolve().parents[2]
    sys.path.insert(0, str(_repo))
    from config.paths import docx_to_md_dirs

    inp, out = docx_to_md_dirs()
    out.mkdir(parents=True, exist_ok=True)
    for p in sorted(inp.glob("*.docx")):
        dest = out / f"{p.stem}_python_docx.md"
        try:
            dest.write_text(docx_to_markdown(p), encoding="utf-8")
            print(f"完成: {p.name} → {dest}")
        except Exception as e:
            print(f"跳过: {p.name} ({e})")


if __name__ == "__main__":
    main()
